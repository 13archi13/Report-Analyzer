import docx
import os

def create_dummy_docx():
    doc = docx.Document()
    doc.add_heading('Technical Report: Properties of Alkanes', 0)
    
    doc.add_paragraph('The following table lists the physical properties of the first five alkanes.')
    
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Alkane'
    hdr_cells[1].text = 'Boiling Point (C)'
    hdr_cells[2].text = 'Melting Point (C)'
    
    data = [
        ('Methane', '-161.5', '-182'),
        ('Ethane', '-89', '-183'),
        ('Propane', '-42', '-188'),
        ('Butane', '-0.5', '-138'),
        ('Pentane', '36.1', '-130')
    ]
    
    for name, bp, mp in data:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = bp
        row_cells[2].text = mp
        
    output_dir = "./data/input"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    filepath = os.path.join(output_dir, "dummy_report.docx")
    doc.save(filepath)
    print(f"Created dummy report at {filepath}")

if __name__ == "__main__":
    create_dummy_docx()
